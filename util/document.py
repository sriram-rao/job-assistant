from __future__ import annotations

import logging
import shutil
import threading
from datetime import datetime
import docx
from pathlib import Path

from docx.document import Document
from docx.text.paragraph import Paragraph

from defaults import RESUME_ORDER
from util import strings


def thread_logger() -> logging.Logger:
    name = f"{__name__}.{threading.current_thread().name}"
    return logging.getLogger(name)


def add_paragraph(text: str, document: Document, index: int, style: str = "Body A", run_style: str = "") -> Paragraph:
    para = document.paragraphs[index + 1].insert_paragraph_before("", style)
    new_run = para.add_run(strings.strip_inline_markdown(text))
    new_run.bold = "bold" in run_style
    new_run.italic = "italic" in run_style
    return para


def set_text(text: str, index: int, document: Document) -> Paragraph:
    document.paragraphs[index].text = strings.strip_inline_markdown(text)
    return document.paragraphs[index]

def add_bullet_point(text: str, document: Document, index: int, style: str = "List Bullet", run_style: str = "") -> Paragraph:
    para = add_paragraph(text, document, index, style, run_style)
    if para._element.pPr is None:
        para._element.get_or_add_pPr()
    para._element.pPr.get_or_add_numPr().get_or_add_numId().val = 1
    para._element.pPr.numPr.get_or_add_ilvl().val = 0
    return para

def add_list(data: list[str], document: Document, index: int, style: str = "List Bullet") -> list[Paragraph]:
    paras: list[Paragraph] = []
    for offset, line in enumerate(data):
        paras.append(add_bullet_point(line, document, index + offset - 1, style, ""))
    return paras

def add_category_list(categories: list[dict[str, list[str]]], document: Document, index: int) -> list[Paragraph]:
    """Add category-based bullet list."""
    lines = []
    for category in categories:
        for key, values in category.items():
            lines.append(f"{key}: {', '.join(strings.strip_inline_markdown(v) for v in values)}")
    return add_list(lines, document, index)

def delete_paragraph(paragraph: Paragraph):
    """Delete a paragraph from the document."""
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def set_list(categories: list[dict[str, list[str]]], index: int, document: Document) -> list[Paragraph]:
    """Replace category list, return number of paragraphs added."""
    paras = add_category_list(categories, document, index)
    # After adding len(paras) paragraphs, the original paragraph at index has shifted to index + len(paras)
    delete_paragraph(document.paragraphs[index + len(paras)])
    return paras


def add_work(data: list[dict[str, str | list[str]]], index: int, document: Document) -> list[Paragraph]:
    """Replace work experience, return net paragraphs added."""
    para = document.paragraphs[index]
    work_offset = 0
    work_paras: list[Paragraph] = []
    for experience in data:
        work_paras.extend(add_work_entry(document, index + work_offset, experience))
        work_offset = len(work_paras) + 1
    delete_paragraph(para)
    return work_paras


def fill_docx(template_path: Path, data: dict[str, object]) -> Document:
    """Replace placeholders in .docx template with formatted content."""
    document = docx.Document(str(template_path))

    offset = 0
    replacements = [
        ("tagline", set_text),  # tagline
        ("languages", set_list),
        ("skills", set_list),
        ("work_experience", add_work),
    ]
    for section, handler in replacements:
        paras = handler(data[section], RESUME_ORDER[section] + offset, document)
        # set_text modifies in place (0 net change), set_list/add_work add N and delete 1 (N-1 net change)
        offset += (len(paras) - 1) if isinstance(paras, list) else 0

    return document


def add_work_entry(doc: Document, index: int, exp: dict[str, str | list[str]]) -> list[Paragraph]:
    """Add single work experience entry with formatting."""
    company_role_line: str = f"{exp.get('company', '')} — {exp.get('role', '')}"
    time_location_line: str = f"{exp.get('start', '')} - {exp.get('end', '')} • " + str(exp.get('location', ''))
    paras: list[Paragraph] = []
    paras.append(add_paragraph(company_role_line, doc, index, style="Body A", run_style="bold"))
    paras.append(add_paragraph(time_location_line, doc, index + len(paras), style="Body A"))
    for bullet in exp.get("bullets", []):
        paras.append(add_bullet_point(str(bullet), doc, index + len(paras)))
    return paras


def archive_old_docx(target_dir: Path) -> None:
    """Move existing .docx and PDF files in target directory to an old/ subdirectory."""
    old_files = list(target_dir.glob("*.docx")) + list(target_dir.glob("*.pdf"))
    if not old_files:
        return

    old_dir = target_dir / "old"
    old_dir.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    for file in old_files:
        old_file = old_dir / f"{timestamp}_{file.name}"
        _ = shutil.move(str(file), str(old_file))

